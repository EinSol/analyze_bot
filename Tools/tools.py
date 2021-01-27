import requests
from decouple import config
from docx import Document
import os
import re
from Tools.requests_url import (summarize_url, extract_article_url, entity_url,
                                sentiment_url, hashtag_url, classify_url)


class Toolkit:
    def __init__(self):
        self.__headers = {
            'x-rapidapi-key': config('API_KEY'),
            'x-rapidapi-host': config('API_HOST')
        }

    def response(self, url, params):
        res = requests.request('GET', url, headers=self.__headers, params=params)
        return res.json()

    def extract_article_info(self, querystring):
        search_url = querystring
        parameters = {'url': search_url}

        return self.response(extract_article_url, parameters)

    def summarize(self, querystring):
        title = querystring.get('title')
        length = querystring.get('length', 5)
        text = querystring.get('text')
        search_url = querystring.get('url')

        parameters = {'title': title,
                      'text': text,
                      'url': search_url,
                      'length': length}

        return self.response(summarize_url, parameters)

    def extract_entity(self, querystring):
        search_url = querystring.get('url')
        text = querystring.get('text')

        parameters = {'url': search_url,
                      'text': text}

        return self.response(entity_url, parameters)

    def extract_sentiment(self, querystring):
        search_url = querystring.get('url')
        text = querystring.get('text')

        parameters = {'url': search_url,
                      'text': text}

        return self.response(sentiment_url, parameters)

    def suggest_hashtag(self, querystring):
        search_url = querystring.get('url')
        text = querystring.get('text')

        parameters = {'url': search_url,
                      'text': text}

        return self.response(hashtag_url, parameters)

    def classify(self, querystring):
        search_url = querystring.get('url')
        text = querystring.get('text')

        parameters = {'url': search_url,
                      'text': text}

        return self.response(classify_url, parameters)

    def create_document(self, querystring, name):
        document = Document()
        document.add_heading('Analyze results', 0)

        for key, value in querystring.items():
            document.add_heading(key, level=1)
            if type(value) == list:
                for subvalue in value:
                    document.add_paragraph(
                        subvalue, style='List Number'
                        )
                continue
            document.add_paragraph(value)

        path = f'{os.getcwd()}/doc_storage/{name}.docx'
        document.save(path)

        return path


def get_parts_of_text(content):
    temp = ''
    text = content.split('.')
    parts = []
    title = text[0]
    for text_part in text:
        if len(temp) > 4900:
            parts.append(temp)
            temp = ''
        temp += re.sub(r'\\x[a-z,0-9]{2}', '', f'{text_part}.'.replace('\\n', ' ')).replace('\\', '')

    if temp:
        parts.append(temp)
    print(len(parts[1]))
    print(parts[1])
    return parts, title
